"""
brd_parser.py
Extracts raw text content from a BRD file (.docx, .pdf, or .txt).
"""
from pathlib import Path


def extract_text(file_path: str) -> str:
    """
    Extract plain text from a BRD document.

    Supports: .docx, .pdf, .txt

    Args:
        file_path: path to the BRD file.

    Returns:
        The extracted text content as a single string.
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"BRD file not found: {file_path}")

    suffix = path.suffix.lower()

    if suffix == ".docx":
        return _extract_docx(path)
    elif suffix == ".pdf":
        return _extract_pdf(path)
    elif suffix == ".txt":
        return path.read_text(encoding="utf-8", errors="ignore")
    else:
        raise ValueError(
            f"Unsupported file type '{suffix}'. Supported: .docx, .pdf, .txt"
        )


def _extract_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts = []

    # Paragraphs (keeps heading text too, since headings are just styled paragraphs)
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    # Tables (BRDs often store requirements in tables)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def _extract_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    parts = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            parts.append(text.strip())
    return "\n".join(parts)
