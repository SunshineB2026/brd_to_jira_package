#!/usr/bin/env python3
"""
brd_to_jira.py

Parses a BRD (docx/pdf/txt/md) and generates a single, consolidated Jira
story (with all acceptance criteria bundled into that one story) using
Claude for the actual writing.

Two modes:

  new     BRD only -> generate a brand new Jira story from scratch.
  update  Existing Jira story + BRD -> generate a detailed/updated version
          of that same story (still one consolidated story).

In both modes the output is always ONE story, not one story per
requirement. Acceptance criteria accumulate inside that single story
even if it gets long.

Usage:
  python brd_to_jira.py new \
      --brd path/to/brd.docx \
      --output out/story.json \
      [--push]

  python brd_to_jira.py update \
      --brd path/to/brd.docx \
      --existing-story path/to/existing_story.json \
      --output out/story.json \
      [--push]

Environment variables:
  ANTHROPIC_API_KEY   required, used to call Claude for generation
  JIRA_BASE_URL        required only if --push is used, e.g. https://yourco.atlassian.net
  JIRA_EMAIL           required only if --push is used
  JIRA_API_TOKEN       required only if --push is used
  JIRA_PROJECT_KEY     required only if --push is used, e.g. "BOX"
"""

import argparse
import json
import os
import sys
from pathlib import Path

try:
    import anthropic
except ImportError:
    anthropic = None

try:
    import docx  # python-docx
except ImportError:
    docx = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    import requests
except ImportError:
    requests = None

try:
    from dotenv import load_dotenv
    load_dotenv()  # picks up ANTHROPIC_API_KEY / JIRA_* from a .env file if present
except ImportError:
    pass


DEFAULT_MODEL = "claude-sonnet-5"


# --------------------------------------------------------------------------
# Input parsing
# --------------------------------------------------------------------------

def extract_text_from_file(path: str) -> str:
    """Extract plain text from a BRD file. Supports .docx, .pdf, .txt, .md."""
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"BRD file not found: {path}")

    suffix = p.suffix.lower()

    if suffix == ".docx":
        if docx is None:
            raise RuntimeError(
                "python-docx is required to read .docx files. "
                "Install with: pip install python-docx"
            )
        document = docx.Document(str(p))
        parts = [para.text for para in document.paragraphs if para.text.strip()]
        # Also pull text out of tables, which BRDs often use for requirements.
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)

    if suffix == ".pdf":
        if pdfplumber is None:
            raise RuntimeError(
                "pdfplumber is required to read .pdf files. "
                "Install with: pip install pdfplumber"
            )
        parts = []
        with pdfplumber.open(str(p)) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    parts.append(text)
        return "\n".join(parts)

    if suffix in (".txt", ".md"):
        return p.read_text(encoding="utf-8")

    raise ValueError(f"Unsupported BRD file type: {suffix}. Use .docx, .pdf, .txt, or .md")


def load_existing_story(path: str) -> dict:
    """
    Load an existing Jira story for the 'update' path.

    Accepts either:
      - a JSON file with fields like key/summary/description/acceptance_criteria
      - a plain .txt/.md file with the story pasted in as free text
    """
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"Existing story file not found: {path}")

    if p.suffix.lower() == ".json":
        with open(p, "r", encoding="utf-8") as f:
            data = json.load(f)
        # Normalize to a consistent shape regardless of what keys were present.
        return {
            "key": data.get("key", ""),
            "summary": data.get("summary", data.get("title", "")),
            "description": data.get("description", ""),
            "acceptance_criteria": data.get("acceptance_criteria", []),
            "positive_test_cases": data.get("positive_test_cases", []),
            "negative_test_cases": data.get("negative_test_cases", []),
            "exit_criteria": data.get("exit_criteria", []),
            "raw": json.dumps(data, indent=2),
        }

    # Fall back to treating it as free text (txt/md/anything else)
    text = p.read_text(encoding="utf-8")
    return {
        "key": "", "summary": "", "description": "",
        "acceptance_criteria": [], "positive_test_cases": [],
        "negative_test_cases": [], "exit_criteria": [], "raw": text,
    }


def _adf_to_text(node) -> str:
    """Flatten a Jira Atlassian Document Format (ADF) body into plain text."""
    if node is None:
        return ""
    if isinstance(node, str):
        return node

    parts = []
    if isinstance(node, dict):
        if node.get("type") == "text":
            return node.get("text", "")
        for child in node.get("content", []) or []:
            parts.append(_adf_to_text(child))
        joined = " ".join(p for p in parts if p)
        # Paragraph-like nodes get their own line.
        if node.get("type") in ("paragraph", "heading", "listItem"):
            return joined + "\n"
        return joined
    if isinstance(node, list):
        return "".join(_adf_to_text(child) for child in node)
    return ""


def fetch_story_from_jira(issue_key: str) -> dict:
    """
    Fetch an existing Jira story live via the Jira REST API, given just its
    issue key (e.g. "BOX-142"), instead of requiring a local export file.
    """
    if requests is None:
        raise RuntimeError("The 'requests' package is required to fetch from Jira.")

    base_url = os.environ.get("JIRA_BASE_URL")
    email = os.environ.get("JIRA_EMAIL")
    token = os.environ.get("JIRA_API_TOKEN")

    missing = [name for name, val in [
        ("JIRA_BASE_URL", base_url),
        ("JIRA_EMAIL", email),
        ("JIRA_API_TOKEN", token),
    ] if not val]
    if missing:
        raise RuntimeError(f"Missing Jira env vars required for --jira-key: {', '.join(missing)}")

    url = f"{base_url.rstrip('/')}/rest/api/3/issue/{issue_key}"
    resp = requests.get(url, auth=(email, token), params={"fields": "summary,description,labels"})
    if not resp.ok:
        raise RuntimeError(f"Jira fetch of {issue_key} failed ({resp.status_code}): {resp.text}")
    data = resp.json()

    fields = data.get("fields", {})
    summary = fields.get("summary", "")
    description = _adf_to_text(fields.get("description")).strip()
    labels = fields.get("labels", [])

    story = {
        "key": data.get("key", issue_key),
        "summary": summary,
        "description": description,
        "acceptance_criteria": [],  # Jira has no dedicated AC field; it's embedded in description text.
        "labels": labels,
    }
    story["raw"] = json.dumps(story, indent=2)
    return story


# --------------------------------------------------------------------------
# Prompting
# --------------------------------------------------------------------------

SYSTEM_PROMPT = """You are a senior business analyst and QA lead who writes Jira stories.

Hard rules you must always follow:
- Produce exactly ONE Jira story. Never split the work into multiple stories,
  sub-tasks, or epics — even if the BRD describes several distinct pieces of
  functionality. Everything belongs in one story.
- All acceptance criteria go inside that single story's acceptance_criteria
  list. It is fine, and expected, for this list to be long if the BRD is
  detailed. Do not compress or drop requirements to keep the story short.
- Write acceptance criteria as clear, testable statements (Given/When/Then
  or plain "should" statements are both fine, be consistent within the
  story).
- Also produce positive test cases (the system behaving correctly for valid,
  expected inputs/flows) and negative test cases (invalid inputs, edge
  cases, error handling, permission failures, boundary conditions). Each
  test case should be a single concrete, testable scenario — specific
  enough that a QA engineer could execute it without guessing, e.g.
  "User submits the form with all required fields filled -> record is
  saved and success message shown" rather than a vague "form works".
- Also produce exit criteria: the conditions that must all be true before
  this story can be considered done and closed (e.g. all acceptance
  criteria met, all positive and negative test cases passed, no open
  blocking bugs, documentation updated, stakeholder sign-off obtained —
  tailor these to what the BRD actually implies, don't pad with generic
  boilerplate that doesn't apply).
- Base everything strictly on the BRD content (and existing story content,
  when provided). Do not invent requirements that aren't supported by the
  input.
- Respond with ONLY a single JSON object, no markdown fences, no preamble,
  no commentary. The JSON object must have exactly these fields:
  {
    "summary": "<one-line Jira story title>",
    "description": "<full story description, background/context and scope>",
    "acceptance_criteria": ["<criterion 1>", "<criterion 2>", "..."],
    "positive_test_cases": ["<test case 1>", "<test case 2>", "..."],
    "negative_test_cases": ["<test case 1>", "<test case 2>", "..."],
    "exit_criteria": ["<condition 1>", "<condition 2>", "..."],
    "labels": ["<optional short labels/tags>"]
  }
"""


def build_new_story_prompt(brd_text: str) -> str:
    return f"""Here is the BRD content:

---BRD START---
{brd_text}
---BRD END---

Generate a brand new Jira story from this BRD, following the rules in your
system prompt (exactly one story, all acceptance criteria, positive/negative
test cases, and exit criteria bundled inside it).
"""


def build_update_story_prompt(brd_text: str, existing_story: dict) -> str:
    return f"""Here is the EXISTING Jira story to update/expand:

---EXISTING STORY START---
{existing_story['raw']}
---EXISTING STORY END---

Here is the BRD content to incorporate:

---BRD START---
{brd_text}
---BRD END---

Produce a detailed, updated version of this SAME story (do not create a
second story). Preserve the intent of the existing story, but expand and
sharpen the description, acceptance criteria, positive/negative test
cases, and exit criteria using the BRD. All acceptance criteria — old and
new — must end up in the single acceptance_criteria list, following the
rules in your system prompt.
"""


# --------------------------------------------------------------------------
# Claude call
# --------------------------------------------------------------------------

def generate_story(prompt: str, model: str = DEFAULT_MODEL) -> dict:
    if anthropic is None:
        raise RuntimeError("The 'anthropic' package is required. Install with: pip install anthropic")

    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        raise RuntimeError("ANTHROPIC_API_KEY environment variable is not set.")

    client = anthropic.Anthropic(api_key=api_key)

    response = client.messages.create(
        model=model,
        max_tokens=16000,
        system=SYSTEM_PROMPT,
        messages=[{"role": "user", "content": prompt}],
    )

    text_blocks = [block.text for block in response.content if block.type == "text"]
    raw_text = "\n".join(text_blocks).strip()

    # Strip accidental code fences if the model adds them despite instructions.
    if raw_text.startswith("```"):
        raw_text = raw_text.strip("`")
        if raw_text.lower().startswith("json"):
            raw_text = raw_text[4:].strip()

    try:
        story = json.loads(raw_text)
    except json.JSONDecodeError as e:
        if response.stop_reason == "max_tokens":
            raise RuntimeError(
                "Claude's response was cut off because it hit the max_tokens limit "
                f"({16000} tokens) before finishing valid JSON — the BRD is large enough "
                "that the full story (description + all acceptance criteria + test cases + "
                "exit criteria) didn't fit. Try increasing max_tokens further in generate_story(), "
                "or trim the BRD to the most relevant sections.\n\n"
                f"Raw (truncated) response was:\n{raw_text}"
            ) from e
        raise RuntimeError(
            f"Claude did not return valid JSON. Raw response was:\n{raw_text}"
        ) from e

    for field in ("summary", "description", "acceptance_criteria"):
        if field not in story:
            raise RuntimeError(f"Claude's response is missing required field '{field}': {story}")

    story.setdefault("positive_test_cases", [])
    story.setdefault("negative_test_cases", [])
    story.setdefault("exit_criteria", [])
    story.setdefault("labels", [])
    return story


# --------------------------------------------------------------------------
# Jira push (optional)
# --------------------------------------------------------------------------

def build_adf_description(story: dict) -> dict:
    """
    Build a proper Atlassian Document Format (ADF) body for the description,
    so it renders correctly in Jira: real paragraphs for the description
    text, a heading, and a real ordered list for acceptance criteria
    (not a wall of text with literal \\n characters, which Jira ignores).
    """
    content = []

    # Description: split on blank lines into separate paragraphs.
    for para in story["description"].split("\n\n"):
        para = para.strip()
        if para:
            content.append({
                "type": "paragraph",
                "content": [{"type": "text", "text": para}],
            })

    def add_list_section(title: str, items: list, ordered: bool = True):
        if not items:
            return
        content.append({
            "type": "heading",
            "attrs": {"level": 3},
            "content": [{"type": "text", "text": title}],
        })
        content.append({
            "type": "orderedList" if ordered else "bulletList",
            **({"attrs": {"order": 1}} if ordered else {}),
            "content": [
                {
                    "type": "listItem",
                    "content": [
                        {"type": "paragraph", "content": [{"type": "text", "text": item}]}
                    ],
                }
                for item in items
            ],
        })

    add_list_section("Acceptance Criteria", story.get("acceptance_criteria", []), ordered=True)
    add_list_section("Positive Test Cases", story.get("positive_test_cases", []), ordered=False)
    add_list_section("Negative Test Cases", story.get("negative_test_cases", []), ordered=False)
    add_list_section("Exit Criteria", story.get("exit_criteria", []), ordered=False)

    return {"type": "doc", "version": 1, "content": content}


def attach_file_to_jira(issue_key: str, file_path: str, base_url: str, email: str, token: str) -> None:
    """Upload a local file as an attachment on an existing Jira issue."""
    p = Path(file_path)
    if not p.exists():
        raise FileNotFoundError(f"File to attach not found: {file_path}")

    url = f"{base_url.rstrip('/')}/rest/api/3/issue/{issue_key}/attachments"
    # Jira requires this header on attachment uploads as an XSRF check.
    headers = {"X-Atlassian-Token": "no-check"}

    with open(p, "rb") as f:
        files = {"file": (p.name, f)}
        resp = requests.post(url, headers=headers, files=files, auth=(email, token))

    if not resp.ok:
        raise RuntimeError(f"Jira attachment upload failed ({resp.status_code}): {resp.text}")


def push_to_jira(story: dict, existing_key: str = "", attach_path: str = "") -> dict:
    if requests is None:
        raise RuntimeError("The 'requests' package is required to push to Jira.")

    base_url = os.environ.get("JIRA_BASE_URL")
    email = os.environ.get("JIRA_EMAIL")
    token = os.environ.get("JIRA_API_TOKEN")
    project_key = os.environ.get("JIRA_PROJECT_KEY")

    missing = [name for name, val in [
        ("JIRA_BASE_URL", base_url),
        ("JIRA_EMAIL", email),
        ("JIRA_API_TOKEN", token),
        ("JIRA_PROJECT_KEY", project_key),
    ] if not val]
    if missing:
        raise RuntimeError(f"Missing Jira env vars required for --push: {', '.join(missing)}")

    description_adf = build_adf_description(story)

    if existing_key:
        # Update existing issue.
        url = f"{base_url.rstrip('/')}/rest/api/3/issue/{existing_key}"
        payload = {
            "fields": {
                "summary": story["summary"],
                "description": description_adf,
            }
        }
        resp = requests.put(url, json=payload, auth=(email, token))
        if not resp.ok:
            raise RuntimeError(f"Jira update failed ({resp.status_code}): {resp.text}")
        if attach_path:
            attach_file_to_jira(existing_key, attach_path, base_url, email, token)
        return {"action": "updated", "key": existing_key}

    # Create new issue.
    url = f"{base_url.rstrip('/')}/rest/api/3/issue"
    payload = {
        "fields": {
            "project": {"key": project_key},
            "summary": story["summary"],
            "description": description_adf,
            "issuetype": {"name": "Story"},
            "labels": story.get("labels", []),
        }
    }
    resp = requests.post(url, json=payload, auth=(email, token))
    if not resp.ok:
        raise RuntimeError(f"Jira create failed ({resp.status_code}): {resp.text}")
    data = resp.json()
    new_key = data.get("key", "")
    if attach_path and new_key:
        attach_file_to_jira(new_key, attach_path, base_url, email, token)
    return {"action": "created", "key": new_key}


# --------------------------------------------------------------------------
# CLI
# --------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(description="Generate a Jira story from a BRD.")
    subparsers = parser.add_subparsers(dest="mode", required=True)

    new_parser = subparsers.add_parser("new", help="Generate a brand new story from a BRD only.")
    new_parser.add_argument("--brd", required=True, help="Path to BRD file (.docx/.pdf/.txt/.md)")
    new_parser.add_argument("--output", required=True, help="Path to write the output JSON story to")
    new_parser.add_argument("--model", default=DEFAULT_MODEL, help="Claude model to use")
    new_parser.add_argument("--push", action="store_true", help="Also push the story to Jira")
    new_parser.add_argument("--attach-brd", action="store_true", help="Attach the BRD file itself to the Jira issue (requires --push)")

    update_parser = subparsers.add_parser("update", help="Enrich an existing story using a BRD.")
    update_parser.add_argument("--brd", required=True, help="Path to BRD file (.docx/.pdf/.txt/.md)")
    story_source = update_parser.add_mutually_exclusive_group(required=True)
    story_source.add_argument("--existing-story", help="Path to existing story (.json/.txt/.md)")
    story_source.add_argument("--jira-key", help="Jira issue key to fetch live, e.g. BOX-142 (requires JIRA_* env vars)")
    update_parser.add_argument("--output", required=True, help="Path to write the output JSON story to")
    update_parser.add_argument("--model", default=DEFAULT_MODEL, help="Claude model to use")
    update_parser.add_argument("--push", action="store_true", help="Also push the updated story to Jira")
    update_parser.add_argument("--attach-brd", action="store_true", help="Attach the BRD file itself to the Jira issue (requires --push)")

    args = parser.parse_args()

    brd_text = extract_text_from_file(args.brd)
    if not brd_text.strip():
        print("Warning: extracted BRD text is empty. Check the input file.", file=sys.stderr)

    existing_key = ""
    if args.mode == "new":
        prompt = build_new_story_prompt(brd_text)
    else:  # update
        if args.jira_key:
            print(f"Fetching {args.jira_key} from Jira...")
            existing_story = fetch_story_from_jira(args.jira_key)
        else:
            existing_story = load_existing_story(args.existing_story)
        existing_key = existing_story.get("key", "")
        prompt = build_update_story_prompt(brd_text, existing_story)

    print(f"Generating story via Claude ({args.model})...")
    story = generate_story(prompt, model=args.model)

    output_path = Path(args.output)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(story, f, indent=2)
    print(f"Story written to {output_path}")
    print(f"  Summary: {story['summary']}")
    print(f"  Acceptance criteria: {len(story['acceptance_criteria'])}")
    print(f"  Positive test cases: {len(story['positive_test_cases'])}")
    print(f"  Negative test cases: {len(story['negative_test_cases'])}")
    print(f"  Exit criteria: {len(story['exit_criteria'])}")

    if args.push:
        print("Pushing to Jira...")
        attach_path = args.brd if args.attach_brd else ""
        result = push_to_jira(story, existing_key=existing_key, attach_path=attach_path)
        print(f"Jira {result['action']} issue: {result['key']}")
        if attach_path:
            print(f"  Attached BRD: {Path(attach_path).name}")


if __name__ == "__main__":
    main()